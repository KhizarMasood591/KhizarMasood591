import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import docx2pdf



class Word():
    def __init__(self, file_path, hotel):
        self.cwd = os.getcwd()
        self.property = hotel
        self.doc = Document(file_path)
        self._riyal_symbol = rf"{self.cwd}/Offer_Letter_Website/Saudi_Riyal_Symbol.png"


    def update_guest_name(self, name):
        para = self.doc.paragraphs[1]
        next_line = self.doc.paragraphs[2]
        next_line.insert_paragraph_before("")
        para.text = f"Dear {name},"
        for run in para.runs:
            run.font.name = 'Lato'
            run.font.size = Pt(10)


    def format_paragraph(self,
                         paragraph,
                         font_name='Lato',
                         font_size=Pt(10),
                         bold=False,
                         font_color=RGBColor(128, 130, 133)):
        paragraph.style.font.name = font_name
        paragraph.style.font.size = font_size
        paragraph.style.font.bold = bold
        paragraph.style.font.color.rgb = font_color


    def update_stay_dates(self, check_in, check_out, los):
        checkin_text = "<checkin>"
        check_out_text = "<checkout>"
        los_text = "<los>"
        paragraphs = self.doc.paragraphs
        for paragraph in paragraphs:
            if checkin_text in paragraph.text:
                paragraph.text = paragraph.text.replace(checkin_text,check_in)
                paragraph.style.font.name = 'Lato'
                paragraph.style.font.size = Pt(10)
                paragraph.style.font.bold = True
                paragraph.style.font.color.rgb = RGBColor(128, 130, 133)
            if check_out_text in paragraph.text:
                paragraph.text = paragraph.text.replace(check_out_text,check_out)
                paragraph.style.font.name = 'Lato'
                paragraph.style.font.size = Pt(10)
                paragraph.style.font.bold = True
                paragraph.style.font.color.rgb = RGBColor(128, 130, 133)
            if los_text in paragraph.text:
                paragraph.text = paragraph.text.replace(los_text,los)
                paragraph.style.font.name = 'Lato'
                paragraph.style.font.size = Pt(10)
                paragraph.style.font.bold = True
                paragraph.style.font.color.rgb = RGBColor(128, 130, 133)



    def update_offer_name(self, offer_name):
        paragraphs = self.doc.paragraphs
        for paragraph in paragraphs:
            if "<Offer Name>" in paragraph.text:
                paragraph.text = paragraph.text.replace("<Offer Name>",offer_name)
                runs = paragraph.runs
                for run in runs:
                    run.underline = True
                    run.font.size = Pt(10)


    def update_rates_table(self, rates: list[dict]):
        table = self.doc.tables[0]
        rows_added = []
        for rate in rates:
            rows_added.append(rate['rowNo'])
            cell_1 = table.cell(rate['rowNo'],2)
            cell_1.text = ""
            for para in cell_1.paragraphs:
                run = para.add_run()
                run.add_picture('Offer_Letter_Website/Saudi_Riyal_Symbol.png',width=Pt(8), height=Pt(8))
                run.add_text(" " + str(rate['roomPerNight']))
            cell_2 = table.cell(rate['rowNo'],3)
            cell_2.text = ""
            for para in cell_2.paragraphs:
                run = para.add_run()
                run.add_picture('Offer_Letter_Website/Saudi_Riyal_Symbol.png', height=Pt(8), width=Pt(8))
                run.add_text(" " + str(rate['roomRate']))
        table_rows = [row_no for row_no in range(1,len(table.rows)) if row_no not in rows_added]
        for row in table_rows:
            range_1 = table.cell(row, 2)
            range_1.text = "No Availability"
            range_2 = table.cell(row, 3)
            range_2.text = "No Availability"
        table.autofit



    def save_files(self, filename, offer_type):
        timestamp = datetime.today().strftime("%Y%m%d-%H%M%S")
        name = offer_type + ' - ' + filename+'_' + timestamp
        docx_path = f'Offer_Letter_Website/word/{name}.docx'
        self.doc.save(docx_path)
        return name